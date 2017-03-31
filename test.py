# -*- coding: UTF-8 -*- 。
import urllib
import urllib2
import re
import time
import os
import sys
import MySQLdb
import hashlib
import json
import random


from docx import Document
from docx.shared import Inches

# -*- coding:utf-8 -*-

# Copyright: Lustralisk
# Author: Cedric Liu
# Date: 2015-11-08

import sys, time

import sys, time

for i in range(5):
    sys.stdout.write(' ' * 10 + '\r')
    sys.stdout.flush()
    sys.stdout.write(str(i) * (5 - i) + '\r')
    sys.stdout.flush()
    time.sleep(1)

class ProgressBar:
    def __init__(self, count=0, total=0, width=50):
        self.count = count
        self.total = total
        self.width = width

    def move(self):
        self.count += 1

    def log(self, s):
        sys.stdout.write(' ' * (self.width + 9) + '\r')
        sys.stdout.flush()
        print s
        progress = self.width * self.count / self.total
        sys.stdout.write('{0:3}/{1:3}: '.format(self.count, self.total))
        sys.stdout.write('#' * progress + '-' * (self.width - progress) + '\r')
        if progress == self.width:
            sys.stdout.write('\n')
        sys.stdout.flush()


bar = ProgressBar(total=10)
for i in range(10):
    bar.move()
    bar.log('We have arrived at: ' + str(i + 1))
    time.sleep(1)

dict = dict.fromkeys(['ja','ja2'], ['大  asd','撒旦    sad'])
dict = dict.get()
print "New Dictionary : %s" %  str(dict)

sys.exit(0)
document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

document.add_picture('monty-truth.png', width=Inches(1.25))

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
# for item in recordset:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(item.qty)
#     row_cells[1].text = str(item.id)
#     row_cells[2].text = item.desc

document.add_page_break()

document.save('demo.docx')

























#http://api.fanyi.baidu.com/api/trans/product/apidoc
#"from":"zh","to":"en",
line = u'参考答案: A'

# gbk -> unicode
# 看需要，这里是把gbk转化为unicode，也可以把utf-8转换为unicode
#line = line.decode('gbk').strip()

# 匹配 '参考答案' 四个字
if re.match(u'^\u53c2\u8003\u7b54\u6848', line):
   q_answer = line
   print q_answer
# f=open('f.txt','a')
# f.write("adw")
# f.write("\t")
# f.write("adw")
print "--------search jp------------"
pattern = re.compile('[\u0800-\u4e00]')
asd = u"　　世界をまたにかける少年の頃に見た写真で忘れられないものの一つに、サマセット？モームの肖像がある。作品集の巻頭の写真で、晩年の作家が、やや右向きに座っている。左手に持つたばこの先から流れる一筋の煙が、"

a = re.findall(u'[\u0800-\u4e00]+',asd.strip())
print len(a)
print "----------search cn-----------"
asd = u"       冬粉撒一三娘方看党阀asfdasd 阿斯达所多  "
a = re.findall(u'[\u0800-\u4e00]+',asd.strip())
print len(a)
if a is not None:
    print a.groups()





print "-------------match-----------------"
asd = u"我叫阿一萨德三扥额阿萨德无,而非"
a = re.match(u'[\u4e00-\u9fa5]+',asd)
print a.group()
sys.exit(0)
conn = MySQLdb.connect(
        host = '127.0.0.1',
        port = 3309,            #Dante 3306  Lab 3309
        user = 'root',
        passwd = '123456',
        db = 'news',
        )
cur = conn.cursor()


import urllib
import urllib2

count = cur.execute('select * from news where news_Url = "/2017/03/27/motorsport/ferrari-f1-hamilton-vettel-australia-battle/index.html"')
info = cur.fetchmany(count)
print info
list = str(info[0][5]).split("  ||||  ")
print info[0][5]
print len(list)
sys.exit(0)
for item,counitem in zip(list,range(1,500)):
    print counitem
    print item
    test_data = {"tgt_text":"apple","src_text":"Apple"}
    test_data['src_text'] = str(item)
    test_data_urlencode = urllib.urlencode(test_data)
    requrl = "http://218.75.34.138:8080/NiuTransServer/translation?from=en&to=zh"
    req = urllib2.Request(url = requrl,data =test_data_urlencode)    
    res_data = urllib2.urlopen(req)
    res = res_data.read()
    print res






# import requests
# import json
#
#
# # POS请求：直接向服务器发送数据
# # get请求：从服务器获取数据
# # 有道，向服务器发送数据，再获取数据
# def get_translate_data(word=None):
#     url = 'http://fanyi.youdao.com/translate?smartresult=dict&smartresult=rule&smartresult=ugc&sessionFrom=null'
#     payload = {'type': 'AUTO', 'i': word, 'doctype': 'json', 'xmlVersion': 1.8,
#                'keyfrom': 'fanyi.web', 'ue': 'UTF-8', 'action': 'FY_BY_CLICKBUTTON',
#                'typoResult': 'true'
#                }  # 建立数据字典
#     response = requests.post(url, data=payload)
#     # print response.text #返回字符串
#
#     content = json.loads(response.text)  # 将字符串转换为json数据
#     print content  # 直接打印，又编码问题，在http://jsoneditoronline.org/中无法查看
#     print json.dumps(content, encoding='utf-8', ensure_ascii=False)  # json，有方法.dumps 实现转码
#
#     print content['translateResult'][0][0]['tgt']
#
#
# if __name__ == '__main__':
#     get_translate_data('苹果')